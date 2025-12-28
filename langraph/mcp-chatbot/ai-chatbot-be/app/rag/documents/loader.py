"""
Document Loader Service
=======================
Production document loading with multiple format support.
Uses LangChain document loaders for consistent handling.
"""

import io
import logging
from typing import List, Optional, Dict, Any
from pathlib import Path

from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    Docx2txtLoader,
    CSVLoader,
    JSONLoader,
    UnstructuredHTMLLoader,
    UnstructuredMarkdownLoader,
)
from langchain_community.document_loaders.base import BaseLoader

# For in-memory loading
import PyPDF2
import docx
import csv
import json

logger = logging.getLogger(__name__)


class DocumentLoaderService:
    """
    Production document loader supporting multiple formats.

    Supported formats:
    - PDF (.pdf)
    - Text (.txt)
    - Word (.docx)
    - CSV (.csv)
    - JSON (.json)
    - HTML (.html)
    - Markdown (.md)
    """

    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "text/plain": "txt",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/csv": "csv",
        "application/json": "json",
        "text/html": "html",
        "text/markdown": "md",
    }

    def __init__(self):
        logger.info("DocumentLoaderService initialized")

    def load_from_bytes(
        self,
        content: bytes,
        content_type: str,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Load document from bytes content.

        Args:
            content: Raw file bytes
            content_type: MIME type of the file
            filename: Original filename
            metadata: Additional metadata to attach

        Returns:
            List of Document objects
        """
        file_type = self.SUPPORTED_TYPES.get(content_type, "unknown")

        if file_type == "unknown":
            logger.warning(f"Unsupported content type: {content_type}")
            # Try to extract as plain text
            return self._load_as_text(content, filename, metadata)

        loader_map = {
            "pdf": self._load_pdf,
            "txt": self._load_text,
            "docx": self._load_docx,
            "csv": self._load_csv,
            "json": self._load_json,
            "html": self._load_html,
            "md": self._load_markdown,
        }

        loader_func = loader_map.get(file_type, self._load_as_text)

        try:
            documents = loader_func(content, filename, metadata)
            logger.info(f"Loaded {len(documents)} document(s) from {filename}")
            return documents
        except Exception as e:
            logger.error(f"Error loading {filename}: {str(e)}")
            raise DocumentLoadError(f"Failed to load {filename}: {str(e)}")

    def load_from_file(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Load document from file path using LangChain loaders.

        Args:
            file_path: Path to the file
            metadata: Additional metadata to attach

        Returns:
            List of Document objects
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        loader_map = {
            ".pdf": PyPDFLoader,
            ".txt": TextLoader,
            ".docx": Docx2txtLoader,
            ".csv": CSVLoader,
            ".html": UnstructuredHTMLLoader,
            ".md": UnstructuredMarkdownLoader,
        }

        loader_class = loader_map.get(extension)

        if not loader_class:
            raise DocumentLoadError(f"Unsupported file type: {extension}")

        try:
            loader = loader_class(file_path)
            documents = loader.load()

            # Add custom metadata
            if metadata:
                for doc in documents:
                    doc.metadata.update(metadata)

            logger.info(f"Loaded {len(documents)} document(s) from {file_path}")
            return documents

        except Exception as e:
            logger.error(f"Error loading {file_path}: {str(e)}")
            raise DocumentLoadError(f"Failed to load {file_path}: {str(e)}")

    def _load_pdf(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load PDF from bytes."""
        documents = []
        pdf_file = io.BytesIO(content)

        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()

                if text.strip():  # Only add non-empty pages
                    doc_metadata = {
                        "source": filename,
                        "page": page_num + 1,
                        "total_pages": len(pdf_reader.pages),
                        "file_type": "pdf",
                    }
                    if metadata:
                        doc_metadata.update(metadata)

                    documents.append(Document(
                        page_content=text,
                        metadata=doc_metadata
                    ))

            # If no text extracted, the PDF might be image-based
            if not documents:
                logger.warning(f"No text extracted from PDF: {filename}. May be image-based.")
                documents.append(Document(
                    page_content="[PDF contains no extractable text - may be image-based]",
                    metadata={"source": filename, "file_type": "pdf", "warning": "no_text_extracted"}
                ))

        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise

        return documents

    def _load_text(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load plain text from bytes."""
        # Try different encodings
        encodings = ["utf-8", "latin-1", "cp1252", "ascii"]
        text = None

        for encoding in encodings:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            raise DocumentLoadError(f"Could not decode {filename} with supported encodings")

        doc_metadata = {
            "source": filename,
            "file_type": "txt",
        }
        if metadata:
            doc_metadata.update(metadata)

        return [Document(page_content=text, metadata=doc_metadata)]

    def _load_docx(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load Word document from bytes."""
        doc_file = io.BytesIO(content)
        doc = docx.Document(doc_file)

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n\n".join(paragraphs)

        doc_metadata = {
            "source": filename,
            "file_type": "docx",
            "paragraph_count": len(paragraphs),
        }
        if metadata:
            doc_metadata.update(metadata)

        return [Document(page_content=text, metadata=doc_metadata)]

    def _load_csv(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load CSV from bytes - each row becomes a document."""
        documents = []
        text = content.decode("utf-8")
        reader = csv.DictReader(io.StringIO(text))

        for row_num, row in enumerate(reader):
            # Convert row to readable text
            row_text = "\n".join(f"{key}: {value}" for key, value in row.items() if value)

            doc_metadata = {
                "source": filename,
                "file_type": "csv",
                "row": row_num + 1,
            }
            if metadata:
                doc_metadata.update(metadata)

            documents.append(Document(
                page_content=row_text,
                metadata=doc_metadata
            ))

        return documents

    def _load_json(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load JSON from bytes."""
        data = json.loads(content.decode("utf-8"))

        # Convert JSON to readable text
        if isinstance(data, list):
            # Each item becomes a document
            documents = []
            for idx, item in enumerate(data):
                text = json.dumps(item, indent=2)
                doc_metadata = {
                    "source": filename,
                    "file_type": "json",
                    "item_index": idx,
                }
                if metadata:
                    doc_metadata.update(metadata)
                documents.append(Document(page_content=text, metadata=doc_metadata))
            return documents
        else:
            # Single document
            text = json.dumps(data, indent=2)
            doc_metadata = {
                "source": filename,
                "file_type": "json",
            }
            if metadata:
                doc_metadata.update(metadata)
            return [Document(page_content=text, metadata=doc_metadata)]

    def _load_html(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load HTML from bytes - strip tags and extract text."""
        from html.parser import HTMLParser

        class HTMLTextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text = []
                self.skip_tags = {"script", "style", "head"}
                self.current_tag = None

            def handle_starttag(self, tag, attrs):
                self.current_tag = tag

            def handle_endtag(self, tag):
                self.current_tag = None

            def handle_data(self, data):
                if self.current_tag not in self.skip_tags:
                    text = data.strip()
                    if text:
                        self.text.append(text)

        html_content = content.decode("utf-8")
        extractor = HTMLTextExtractor()
        extractor.feed(html_content)

        text = "\n".join(extractor.text)

        doc_metadata = {
            "source": filename,
            "file_type": "html",
        }
        if metadata:
            doc_metadata.update(metadata)

        return [Document(page_content=text, metadata=doc_metadata)]

    def _load_markdown(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Load Markdown from bytes."""
        text = content.decode("utf-8")

        doc_metadata = {
            "source": filename,
            "file_type": "markdown",
        }
        if metadata:
            doc_metadata.update(metadata)

        return [Document(page_content=text, metadata=doc_metadata)]

    def _load_as_text(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """Fallback: try to load as plain text."""
        return self._load_text(content, filename, metadata)

    @staticmethod
    def is_supported(content_type: str) -> bool:
        """Check if content type is supported."""
        return content_type in DocumentLoaderService.SUPPORTED_TYPES


class DocumentLoadError(Exception):
    """Exception raised when document loading fails."""
    pass


# Global instance
document_loader = DocumentLoaderService()
